"""DOCX generation service."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import List
from io import BytesIO
from app.services.pdf_processor import PDFPage, PDFTextBlock, PDFImage
from app.services.font_mapper import FontMapper


class DOCXGenerator:
    """Handles DOCX file generation from processed PDF data."""
    
    def __init__(self):
        """Initialize DOCX generator."""
        self.font_mapper = FontMapper()
    
    def create_document(self, pages: List[PDFPage], output_path: Path) -> Path:
        """
        Create DOCX document from processed PDF pages.
        
        Args:
            pages: List of processed PDF pages
            output_path: Path to save DOCX file
            
        Returns:
            Path to generated DOCX file
        """
        doc = Document()
        
        for page in pages:
            self._add_page_to_document(doc, page)
            
            # Add page break (except for last page)
            if page.page_number < len(pages) - 1:
                doc.add_page_break()
        
        doc.save(output_path)
        return output_path
    
    def _add_page_to_document(self, doc: Document, page: PDFPage) -> None:
        """
        Add a single PDF page content to Word document.
        
        Args:
            doc: python-docx Document object
            page: Processed PDF page
        """
        # Group text blocks into paragraphs
        # Simple approach: group by vertical position (y0)
        paragraphs = self._group_text_into_paragraphs(page.text_blocks)
        
        for para_blocks in paragraphs:
            p = doc.add_paragraph()
            
            for block in para_blocks:
                # Map font
                mapped_font = self.font_mapper.map_font(block.font_name)
                
                # Add run with formatting
                run = p.add_run(block.text + " ")
                run.font.name = mapped_font
                run.font.size = Pt(block.font_size)
        
        # Add images
        for img in page.images:
            self._add_image_to_document(doc, img)
    
    def _group_text_into_paragraphs(
        self, 
        text_blocks: List[PDFTextBlock],
        y_tolerance: float = 5.0
    ) -> List[List[PDFTextBlock]]:
        """
        Group text blocks into paragraphs based on vertical position.
        
        Args:
            text_blocks: List of text blocks
            y_tolerance: Vertical tolerance for grouping (in points)
            
        Returns:
            List of paragraph groups
        """
        if not text_blocks:
            return []
        
        # Sort by vertical position
        sorted_blocks = sorted(text_blocks, key=lambda b: (b.y0, b.x0))
        
        paragraphs = []
        current_para = [sorted_blocks[0]]
        current_y = sorted_blocks[0].y0
        
        for block in sorted_blocks[1:]:
            # If vertical position is similar, add to current paragraph
            if abs(block.y0 - current_y) <= y_tolerance:
                current_para.append(block)
            else:
                # Start new paragraph
                paragraphs.append(current_para)
                current_para = [block]
                current_y = block.y0
        
        # Add last paragraph
        if current_para:
            paragraphs.append(current_para)
        
        return paragraphs
    
    def _add_image_to_document(self, doc: Document, img: PDFImage) -> None:
        """
        Add image to Word document.
        
        Args:
            doc: python-docx Document object
            img: PDF image data
        """
        try:
            # Convert bytes to BytesIO
            image_stream = BytesIO(img.data)
            
            # Calculate width (preserve aspect ratio, max 6 inches)
            max_width = Inches(6)
            aspect_ratio = img.height / img.width if img.width > 0 else 1
            width = min(max_width, Inches(img.width / 72))  # Convert pixels to inches
            
            # Add image
            doc.add_picture(image_stream, width=width)
            
            # Add spacing after image
            doc.add_paragraph()
            
        except Exception as e:
            # If image fails to add, log error and skip
            print(f"Warning: Failed to add image: {e}")
